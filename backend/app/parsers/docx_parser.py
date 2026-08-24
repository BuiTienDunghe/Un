from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.parsers.base_parser import BaseParser


# T17: Word marks a heading with a paragraph STYLE, not with the "#" prefix the
# chunker looks for, so styled headings used to reach chunking as plain text and
# every .docx chunk came out with heading_path = NULL.  Both the display name
# ("Heading 2") and the style id ("Heading2") are checked because the name is
# localized in non-English Word installs while the id usually is not.
_HEADING_LEVEL = re.compile(r"heading\s*([1-9])", re.IGNORECASE)
_TITLE_STYLES = {"title", "subtitle"}
_MAX_HEADING_LEVEL = 6  # chunking._HEADING_PATTERN accepts #{1,6}

_TXBX_CONTENT = qn("w:txbxContent")
# python-docx's nsmap has no "mc" prefix, so the markup-compatibility namespace
# is spelled out here.
_MC_FALLBACK = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"


class DocxParser(BaseParser):
    """Read a .docx the way the chunker expects to be fed.

    `document.paragraphs` alone silently drops every table, every text box and
    every heading level: measured on three real documents it recovered 79-92%
    of the body text, and the missing part was concentrated in exactly the
    high-value content (results tables, session summary boxes).  This walks the
    body in document order instead, renders tables as markdown pipe rows that
    chunking._is_table_line recognizes, and restores heading markers.

    Pages stay None: a .docx has no fixed pagination, and none of the sample
    documents carried even a rendered page break to approximate one.
    """

    def parse(self, path: Path) -> list[tuple[int | None, str]]:
        document = Document(path)
        blocks: list[str] = []
        for item in _iter_body(document):
            if isinstance(item, Paragraph):
                blocks.extend(_paragraph_blocks(item))
            else:
                blocks.extend(_table_blocks(item))
        return [(None, "\n\n".join(blocks))]


def _iter_body(parent: DocxDocument | _Cell):
    """Yield Paragraph and Table children in document order.

    python-docx exposes `.paragraphs` and `.tables` as two separate flat lists,
    which loses both the interleaving and any table nested inside a cell.
    """
    element = parent.element.body if isinstance(parent, DocxDocument) else parent._tc
    for child in element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _heading_prefix(paragraph: Paragraph) -> str:
    style = paragraph.style
    name, style_id = (style.name or ""), (style.style_id or "")
    match = _HEADING_LEVEL.search(name) or _HEADING_LEVEL.search(style_id)
    if match:
        return "#" * min(int(match.group(1)), _MAX_HEADING_LEVEL) + " "
    return "# " if name.strip().lower() in _TITLE_STYLES else ""


def _paragraph_blocks(paragraph: Paragraph) -> list[str]:
    blocks: list[str] = []
    text = paragraph.text.strip()
    if text:
        # A heading spanning several lines would break the "#" convention, and
        # the chunker treats one heading as one line.
        blocks.append(_heading_prefix(paragraph) + " ".join(text.split()))
    blocks.extend(_textbox_blocks(paragraph))
    return blocks


def _textbox_blocks(paragraph: Paragraph) -> list[str]:
    """Text inside shapes/text boxes, which `Paragraph.text` never returns.

    Word wraps a modern shape in mc:AlternateContent and writes the SAME text
    twice: once under mc:Choice and once as a legacy VML copy under
    mc:Fallback. Taking both would duplicate the content, so the fallback copy
    is skipped.
    """
    blocks: list[str] = []
    for container in paragraph._p.iter(_TXBX_CONTENT):
        if any(ancestor.tag == _MC_FALLBACK for ancestor in container.iterancestors()):
            continue
        lines = [" ".join(node.text.split()) for node in container.iter(qn("w:t")) if node.text and node.text.strip()]
        if lines:
            blocks.append(" ".join(lines))
    return blocks


def _cell_text(cell: _Cell) -> str:
    """One table cell flattened to a single pipe-safe line."""
    parts: list[str] = []
    for item in _iter_body(cell):
        if isinstance(item, Paragraph):
            parts.extend(_paragraph_blocks(item))
        else:  # a table nested inside this cell
            parts.extend(_table_blocks(item))
    # Newlines would end the table block early and "|" would forge a column.
    return " ".join(" ".join(parts).split()).replace("|", "\\|")


def _table_blocks(table: Table) -> list[str]:
    """Render a table as markdown pipe rows.

    chunking._is_table_line only accepts a line starting with "|" and holding
    at least two of them, and chunking._table_blocks repeats the header row on
    every split when the second line is a "---" separator — so emitting that
    separator is what keeps a long table readable after chunking.
    """
    rows = [[_cell_text(cell) for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return []
    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join("---" for _ in rows[0]) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return ["\n".join(lines)]
