"""T17: the .docx reader must hand the chunker what the chunker looks for.

Fixtures are built with python-docx rather than committed as binaries: the
repository is public, and a generated document states the structure under test
in the test itself.
"""
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml

from app.parsers.docx_parser import DocxParser
from app.utils.chunking import chunk_pages


W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
MC = 'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'


def _build(path: Path) -> Path:
    document = Document()
    document.add_heading("Guide", level=1)
    document.add_paragraph("Opening paragraph of the guide body.")
    document.add_heading("Setup", level=2)
    document.add_paragraph("Install the tool before anything else.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Model"
    table.cell(0, 1).text = "BLEU"
    table.cell(1, 0).text = "Transformer"
    table.cell(1, 1).text = "41.8"
    document.add_paragraph("Closing paragraph after the table.")
    document.save(path)
    return path


def test_docx_keeps_document_order_headings_and_tables(tmp_path: Path):
    text = DocxParser().parse(_build(tmp_path / "doc.docx"))[0][1]

    # Heading styles become the "#" markers chunking._HEADING_PATTERN needs.
    assert "# Guide" in text and "## Setup" in text
    # The table survives at all -- document.paragraphs alone dropped it.
    assert "| Model | BLEU |" in text and "| Transformer | 41.8 |" in text
    # And it is a markdown table, so chunking._table_blocks can repeat the header.
    assert "| --- | --- |" in text
    # Document order: the table sits between its two surrounding paragraphs.
    assert text.index("Install the tool") < text.index("| Model") < text.index("Closing paragraph")


def test_docx_chunks_carry_heading_path_and_table_block_type(tmp_path: Path):
    pages = DocxParser().parse(_build(tmp_path / "doc.docx"))
    # Small budget, no overlap: _make_chunk only labels a chunk whose blocks all
    # share one heading path, and overlap would carry the preceding paragraph
    # into the table chunk and make it "mixed".
    chunks = chunk_pages([(page, body, "native") for page, body in pages], 12, 0)

    assert any(chunk.heading_path == ("Guide", "Setup") for chunk in chunks)
    table = [chunk for chunk in chunks if chunk.block_type == "table"]
    assert len(table) == 1 and "| Transformer | 41.8 |" in table[0].content
    assert table[0].heading_path == ("Guide", "Setup")
    # A .docx has no fixed pagination and Word writes no page break here.
    assert all(chunk.page_start is None for chunk in chunks)


def test_docx_cell_text_cannot_forge_a_column_or_end_the_table(tmp_path: Path):
    path = tmp_path / "pipe.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header"
    table.cell(0, 1).text = "Other"
    table.cell(1, 0).text = "a | b"          # a literal pipe
    table.cell(1, 1).text = "line1\nline2"   # a newline inside the cell
    document.save(path)

    text = DocxParser().parse(path)[0][1]

    assert r"a \| b" in text
    assert "| line1 line2 |" in text
    assert all(line.startswith("|") for line in text.splitlines() if line.strip())


def test_docx_textbox_is_read_once_not_twice(tmp_path: Path):
    """Word stores a shape twice: mc:Choice and a legacy mc:Fallback copy."""
    path = tmp_path / "textbox.docx"
    document = Document()
    paragraph = document.add_paragraph("Body paragraph.")
    boxed = (
        f'<w:r {W} {MC}><mc:AlternateContent>'
        f'<mc:Choice Requires="wps"><w:txbxContent><w:p><w:r><w:t>BOXED SUMMARY</w:t></w:r></w:p></w:txbxContent></mc:Choice>'
        f'<mc:Fallback><w:txbxContent><w:p><w:r><w:t>BOXED SUMMARY</w:t></w:r></w:p></w:txbxContent></mc:Fallback>'
        f'</mc:AlternateContent></w:r>'
    )
    paragraph._p.append(parse_xml(boxed))
    document.save(path)

    text = DocxParser().parse(path)[0][1]

    assert text.count("BOXED SUMMARY") == 1
    assert "Body paragraph." in text
